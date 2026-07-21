from io import BytesIO

import fitz
from fastapi.testclient import TestClient

from qadam_api.api.dependencies import get_orchestrator


def _short_pdf() -> bytes:
    document = fitz.open()
    document.new_page().insert_text((72, 72), "Signature")
    buffer = document.tobytes()
    document.close()
    return buffer


def test_upload_returns_accepted_analysis_and_private_access_token(client: TestClient) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("Договор аренды квартиры")
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/v1/analyses",
        files={"file": ("contract.docx", buffer.getvalue())},
    )

    assert response.status_code == 202
    assert response.json()["status"] == "queued"
    assert response.json()["analysis_id"]
    assert len(response.json()["access_token"]) >= 32
    assert "filename" not in response.json()


def test_polling_and_findings_require_matching_token(
    client: TestClient, uploaded_analysis: tuple[object, str]
) -> None:
    analysis_id, token = uploaded_analysis

    unauthorized = client.get(f"/api/v1/analyses/{analysis_id}")
    wrong = client.get(f"/api/v1/analyses/{analysis_id}", headers={"X-Analysis-Token": "wrong"})
    report = client.get(f"/api/v1/analyses/{analysis_id}", headers={"X-Analysis-Token": token})
    findings = client.get(
        f"/api/v1/analyses/{analysis_id}/findings",
        headers={"X-Analysis-Token": token},
    )

    assert unauthorized.status_code == 401
    assert wrong.status_code == 404
    assert wrong.headers["content-type"].startswith("application/problem+json")
    assert report.status_code == 200
    assert report.json()["severity_counts"]["high"] == 1
    assert findings.status_code == 200
    assert findings.json()[0]["citations"][0]["source_id"] == "civil-lease-552"


def test_polling_response_does_not_expose_private_document_evidence(
    client: TestClient, uploaded_analysis: tuple[object, str]
) -> None:
    analysis_id, token = uploaded_analysis

    response = client.get(
        f"/api/v1/analyses/{analysis_id}",
        headers={"X-Analysis-Token": token},
    )

    assert response.status_code == 200
    assert "evidence" not in response.json()
    assert response.json()["question_suggestions"] == ["Как изменить условие о депозите?"]
    assert "Караганда" not in response.text


def test_unsupported_upload_uses_safe_problem_details(client: TestClient) -> None:
    response = client.post(
        "/api/v1/analyses",
        files={"file": ("notes.txt", b"not a contract", "text/plain")},
    )

    assert response.status_code == 415
    assert response.headers["content-type"].startswith("application/problem+json")
    assert response.json()["code"] == "unsupported_type"
    assert "path" not in response.text.casefold()


def test_ocr_required_is_a_recoverable_analysis_failure(client: TestClient) -> None:
    response = client.post(
        "/api/v1/analyses",
        files={"file": ("scan.pdf", _short_pdf(), "application/pdf")},
    )
    payload = response.json()
    report = client.get(
        f"/api/v1/analyses/{payload['analysis_id']}",
        headers={"X-Analysis-Token": payload["access_token"]},
    )

    assert response.status_code == 202
    assert payload["status"] == "queued"
    assert report.json()["failure_code"] == "ocr_required"


def test_browser_origin_can_preflight_contract_upload(client: TestClient) -> None:
    response = client.options(
        "/api/v1/analyses",
        headers={
            "Origin": "http://localhost:3000",
            "Access-Control-Request-Method": "POST",
            "Access-Control-Request-Headers": "content-type,x-analysis-token",
        },
    )

    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://localhost:3000"
    assert "x-analysis-token" in response.headers["access-control-allow-headers"].casefold()


def test_unexpected_pipeline_failure_becomes_safe_pollable_state(client: TestClient) -> None:
    from docx import Document

    class RaisingOrchestrator:
        def analyze(self, upload: object, *, on_stage: object = None) -> object:
            raise RuntimeError("/private/server/path must never escape")

    client.app.dependency_overrides[get_orchestrator] = RaisingOrchestrator
    document = Document()
    document.add_paragraph("Договор аренды квартиры")
    buffer = BytesIO()
    document.save(buffer)

    accepted = client.post(
        "/api/v1/analyses",
        files={"file": ("contract.docx", buffer.getvalue())},
    )
    payload = accepted.json()
    report = client.get(
        f"/api/v1/analyses/{payload['analysis_id']}",
        headers={"X-Analysis-Token": payload["access_token"]},
    )

    assert accepted.status_code == 202
    assert report.json()["status"] == "failed"
    assert report.json()["failure_code"] == "internal_error"
    assert "/private/server/path" not in report.text
