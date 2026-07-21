from io import BytesIO

import fitz
import pytest
from docx import Document

from qadam_api.documents.parser import DocumentParsingError, parse_document
from qadam_api.documents.validation import validate_upload


def make_pdf_bytes(*pages: str) -> bytes:
    document = fitz.open()
    for text in pages:
        page = document.new_page()
        page.insert_textbox((72, 72, 520, 760), text, fontsize=11)
    result = document.tobytes()
    document.close()
    return result


def make_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Договор найма жилого помещения", level=1)
    document.add_paragraph(
        "Арендодатель передает, а арендатор принимает квартиру во временное пользование."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Арендная плата"
    table.cell(0, 1).text = "180 000 тенге"
    table.cell(1, 0).text = "Срок оплаты"
    table.cell(1, 1).text = "до пятого числа каждого месяца"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parses_pdf_into_normalized_blocks_with_page_spans() -> None:
    content = make_pdf_bytes(
        "Residential lease agreement between the landlord and the tenant. "
        "The apartment is transferred for temporary use at the stated monthly rent.",
        "The tenant pays utilities separately and returns the apartment under a handover act. "
        "Both parties provide written notice before termination.",
    )
    upload = validate_upload(filename="agreement.pdf", content=content)

    parsed = parse_document(upload)

    assert [block.page_number for block in parsed.blocks] == [1, 2]
    assert parsed.blocks[0].index == 0
    assert parsed.blocks[0].kind == "page_text"
    assert "monthly rent" in parsed.full_text
    assert "written notice" in parsed.full_text


def test_parses_docx_paragraphs_and_table_rows() -> None:
    upload = validate_upload(filename="agreement.docx", content=make_docx_bytes())

    parsed = parse_document(upload)

    assert any(block.kind == "paragraph" for block in parsed.blocks)
    assert any(block.kind == "table_row" for block in parsed.blocks)
    assert any("180 000 тенге" in block.text for block in parsed.blocks)
    assert [block.index for block in parsed.blocks] == list(range(len(parsed.blocks)))
    assert "Договор найма жилого помещения" in parsed.full_text


def test_normalizes_repeated_whitespace_without_joining_blocks() -> None:
    content = make_pdf_bytes(
        "Residential    lease agreement with enough text to pass extraction.\n\n"
        "Monthly rent is 180000 tenge and utilities are paid separately by the tenant."
    )
    upload = validate_upload(filename="agreement.pdf", content=content)

    parsed = parse_document(upload)

    assert "Residential lease agreement" in parsed.full_text
    assert "    " not in parsed.full_text


def test_marks_document_with_too_little_text_as_ocr_required() -> None:
    upload = validate_upload(filename="scan.pdf", content=make_pdf_bytes("Signature"))

    with pytest.raises(DocumentParsingError) as error:
        parse_document(upload)

    assert error.value.code == "ocr_required"
