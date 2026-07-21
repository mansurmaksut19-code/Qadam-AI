"""Extract ordered, source-addressable text from validated uploads."""

from __future__ import annotations

import re
from io import BytesIO

import fitz
from docx import Document

from qadam_api.documents.types import ParsedDocument, TextBlock, ValidatedUpload
from qadam_api.documents.validation import DOCX_MEDIA_TYPE, PDF_MEDIA_TYPE

MIN_USEFUL_CHARACTERS = 80
_WHITESPACE = re.compile(r"[\t\r\f\v ]+")


class DocumentParsingError(ValueError):
    """Stable parsing failure suitable for API problem details."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def _normalize_text(value: str) -> str:
    lines = (_WHITESPACE.sub(" ", line).strip() for line in value.replace("\xa0", " ").split("\n"))
    return "\n".join(line for line in lines if line)


def _parse_pdf(upload: ValidatedUpload) -> list[TextBlock]:
    document = fitz.open(stream=upload.content, filetype="pdf")
    try:
        blocks: list[TextBlock] = []
        for page_index, page in enumerate(document):
            text = _normalize_text(page.get_text("text"))
            if text:
                blocks.append(
                    TextBlock(
                        index=len(blocks),
                        text=text,
                        kind="page_text",
                        page_number=page_index + 1,
                    )
                )
        return blocks
    finally:
        document.close()


def _parse_docx(upload: ValidatedUpload) -> list[TextBlock]:
    document = Document(BytesIO(upload.content))
    blocks: list[TextBlock] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = _normalize_text(paragraph.text)
        if text:
            blocks.append(
                TextBlock(
                    index=len(blocks),
                    text=text,
                    kind="paragraph",
                    paragraph_number=paragraph_index,
                )
            )

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                blocks.append(
                    TextBlock(
                        index=len(blocks),
                        text=text,
                        kind="table_row",
                    )
                )
    return blocks


def parse_document(upload: ValidatedUpload) -> ParsedDocument:
    """Extract normalized blocks or explicitly require OCR for image-only content."""

    if upload.media_type == PDF_MEDIA_TYPE:
        blocks = _parse_pdf(upload)
    elif upload.media_type == DOCX_MEDIA_TYPE:
        blocks = _parse_docx(upload)
    else:  # pragma: no cover - ValidatedUpload is only created by the validator.
        raise DocumentParsingError("unsupported_type", "Неподдерживаемый формат документа.")

    full_text = "\n\n".join(block.text for block in blocks)
    useful_characters = sum(character.isalnum() for character in full_text)
    if useful_characters < MIN_USEFUL_CHARACTERS:
        raise DocumentParsingError(
            "ocr_required",
            "В документе недостаточно извлекаемого текста. Нужна OCR-версия или DOCX.",
        )

    return ParsedDocument(blocks=tuple(blocks), full_text=full_text)
